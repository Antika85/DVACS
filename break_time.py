from docx import Document
from tkinter import messagebox
from paragraps_and_sentences_script import fun_find_paragraphs_and_sentences

def break_time_fun(file):
    try:
        paragraphs_and_sentences, all_sentences = fun_find_paragraphs_and_sentences(file)

        max_chars = 1950
        current_chars = 0
        file_count = 1
        index = 1
        time = 3
        doc = Document()

        for p in all_sentences:
            if current_chars + len(p) > max_chars:
                doc.save(f'document_{file_count}.docx')
                file_count += 1
                index = 1  # Сбрасываем индекс для нового файла
                current_chars = 0  # Обнуляем текущий счетчик символов
                doc = Document()

            doc.add_paragraph(p)
            pause_paragraph = doc.add_paragraph()
            pause_paragraph.add_run(f'<b class="tag-ssml param-ssml ssml_single ssml-type_pause" '
                                                f'id="ssml_range_{index}" contenteditable="false" '
                                                f'data-tag-id="{index}" data-tag-type="pause" '
                                                f'data-tag-value="{time}s">\n<span class="ssml-value">{time}s</span></b>')
            current_chars += len(p)
            current_chars += 2
            index += 1

        doc.save(f'document_{file_count}.docx')
        messagebox.showinfo("Успех!", f"Текстовый(ые) документ(ы) успешно созданы!")

    except Exception as ex:
        messagebox.showerror("Ошибка", f"Произошла неизвестная ошибка, покажите её автору: {ex}")
